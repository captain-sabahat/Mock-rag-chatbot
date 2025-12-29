"""
================================================================================
FILE: src/utils.py
================================================================================

PURPOSE:
Shared utility functions used across the backend. Includes cache key
generation, common helpers, basic text processing, and **file handling utilities**.
Enables code reuse and consistency.

WORKFLOW:
1. SECTION 1: Cache utilities - Generate consistent cache keys
2. SECTION 2: Common helpers - Request ID, logging, retry logic
3. SECTION 3: Text utilities (basic) - Text cleaning, truncation
4. **SECTION 4: File utilities (NEW) - File parsing, text extraction**

IMPORTS:
- hashlib: Hash functions for cache keys
- uuid: Unique ID generation
- logging: Logging utilities
- typing: Type hints
- **io: BytesIO for in-memory file operations** (NEW)
- **mimetypes: File type detection** (NEW)
- **re: Text processing for file content** (NEW)

INPUTS:
- Various (user_id, query, context data, file content, etc.)

OUTPUTS:
- Cache keys (str), request IDs (str), processed text (str)
- **Extracted text from files (str)** (NEW)
- **File metadata (dict)** (NEW)

KEY FACTS:
- No imports from src modules (prevents circular dependencies)
- All functions are pure/stateless (idempotent)
- Cache key format is consistent across all uses
- Request ID enables request tracking/correlation
- **All file operations in-memory (no disk writes)** (NEW)
- **Thread-safe for concurrent file processing** (NEW)

CACHE KEY STRATEGY:
- Namespace-based keys: prefix:type:identifier
- Examples:
  - Query cache: user:USER_ID:query:QUERY_HASH
  - Session cache: session:USER_ID:SESSION_ID
  - Document summary: summary:USER_ID:DOC_HASH
  - **Ephemeral doc: ephemeral:USER_ID:DOC_HASH** (NEW)
- Hash queries to keep key length reasonable

FILE HANDLING STRATEGY (NEW):
- No temporary files on disk
- Content streamed into BytesIO
- File extension determines parsing method
- PDF: Binary content (sent to embedding model)
- TXT/MD: Direct text extraction
- DOCX: XML parsing for text extraction
- All text normalized before processing

FUTURE SCOPE (Phase 2+):
- Add more text utilities (tokenization, PII masking)
- Add time utilities (TTL calculation, rate limiting)
- Add retry utilities (exponential backoff)
- Add async utilities (batch processing)
- Add encryption utilities (secure data)
- Add validation utilities (regex, format checking)
- Add encoding utilities (base64, json)
- **Add OCR support (Tesseract)** (Phase 2)
- **Add virus scanning (ClamAV)** (Phase 2)
- **Add advanced PDF parsing (pdfplumber)** (Phase 2)
- **Add DOCX table extraction** (Phase 2)

TESTING ENVIRONMENT:
- Mock cache keys in tests
- Verify hash consistency
- Test edge cases (empty strings, special chars)
- Verify request ID uniqueness
- **Create mock file objects for testing** (NEW)
- **Test text extraction with sample files** (NEW)
- **Test file size handling** (NEW)

PRODUCTION DEPLOYMENT:
- Cache keys must be consistent across servers
- Request IDs used for correlation logging
- Hash functions deterministic (same input = same output)
- **File parsing must handle malformed files gracefully** (NEW)
- **Memory limits enforced for file operations** (NEW)
- **Timeout limits for file processing** (NEW)
"""
#================================================================================
#IMPORTS
#================================================================================

import hashlib
import uuid
import logging
from typing import Optional, Dict, Any
from io import BytesIO
import mimetypes
import re
import asyncio

logger = logging.getLogger(__name__)

# ============================================================================
# SECTION 1: CACHE UTILITIES
# ============================================================================

def generate_query_cache_key(user_id: str, query_hash: str) -> str:
    """Generate cache key for user query result."""
    return f"user:{user_id}:query:{query_hash}"


def generate_session_cache_key(user_id: str, session_id: str) -> str:
    """Generate cache key for user session data."""
    return f"session:{user_id}:{session_id}"


def generate_document_summary_cache_key(user_id: str, doc_hash: str) -> str:
    """Generate cache key for document summary."""
    return f"summary:{user_id}:{doc_hash}"


def generate_ephemeral_doc_cache_key(
    user_id: str, doc_hash: str, session_id: str
) -> str:
    """
    Generate cache key for ephemeral document (temporary session-scoped).
    
    Format: ephemeral:USER_ID:SESSION_ID:DOC_HASH
    
    NOTE: Ephemeral docs typically NOT cached - this is for future use.
    """
    return f"ephemeral:{user_id}:{session_id}:{doc_hash}"


def hash_query(query: str) -> str:
    """Generate stable hash of query (for cache key)."""
    return hashlib.sha256(query.encode()).hexdigest()


def hash_document(content: bytes) -> str:
    """Generate stable hash of document (for cache key)."""
    return hashlib.sha256(content).hexdigest()


def calculate_ttl(base_ttl_seconds: int, priority: str = "normal") -> int:
    """Calculate TTL based on priority."""
    multipliers = {
        "high": 2.0,      # Cache 2x longer
        "normal": 1.0,    # Base TTL
        "low": 0.5        # Cache 0.5x (shorter)
    }
    multiplier = multipliers.get(priority, 1.0)
    return int(base_ttl_seconds * multiplier)


# ============================================================================
# SECTION 2: COMMON HELPERS
# ============================================================================

def generate_request_id() -> str:
    """Generate unique request ID (UUID v4)."""
    return str(uuid.uuid4())


def format_logger_context(
    request_id: str,
    user_id: str,
    session_id: Optional[str] = None
) -> Dict[str, str]:
    """Format structured logging context."""
    context = {
        "request_id": request_id,
        "user_id": user_id
    }
    if session_id:
        context["session_id"] = session_id
    return context


def format_error_context(error_code: str, error_message: str, **kwargs) -> Dict[str, Any]:
    """Format error context for logging."""
    return {
        "error_code": error_code,
        "error_message": error_message,
        "context": kwargs
    }


async def retry_async_with_backoff(
    func,
    max_retries: int = 3,
    backoff_factor: float = 2.0,
    *args,
    **kwargs
):
    """Retry async function with exponential backoff."""
    last_exception = None
    
    for attempt in range(max_retries):
        try:
            return await func(*args, **kwargs)
        except Exception as e:
            last_exception = e
            if attempt < max_retries - 1:
                wait_time = (backoff_factor ** attempt)
                logger.warning(
                    f"Retry {attempt + 1}/{max_retries} after {wait_time}s: {str(e)}"
                )
                await asyncio.sleep(wait_time)
    
    # All retries exhausted
    raise last_exception


# ============================================================================
# SECTION 3: TEXT UTILITIES
# ============================================================================

def truncate_text(text: str, max_length: int) -> str:
    """Truncate text to maximum length."""
    if len(text) > max_length:
        return text[:max_length]
    return text


def clean_text(text: str) -> str:
    """Clean text (whitespace normalization)."""
    # Normalize whitespace (multiple spaces → single space)
    return " ".join(text.split())


def sanitize_text(text: str, remove_special_chars: bool = False) -> str:
    """Sanitize text for safe usage."""
    # Basic sanitization
    text = clean_text(text)
    if remove_special_chars:
        # Remove non-alphanumeric (keep spaces, basic punctuation)
        text = re.sub(r'[^a-zA-Z0-9\s\.\,\!\?]', '', text)
    return text


def extract_keywords(text: str, num_keywords: int = 5) -> list[str]:
    """Extract keywords from text (basic implementation)."""
    # Simple implementation: split by spaces, remove stop words
    stop_words = {
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 
        'to', 'for', 'of', 'is', 'are', 'was', 'be'
    }
    
    words = text.lower().split()
    keywords = [w for w in words if w not in stop_words and len(w) > 3]
    return keywords[:num_keywords]


def split_text(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 100,
) -> list[str]:
    """
    Split long text into overlapping chunks.
    
    IMPORTANT: This is for text that exceeds SLM input limits ONLY.
    For ephemeral documents: 
    - Do NOT use split_text for chunking before embedding
    - Use split_text only if document > SLM max input size
    - Pass each chunk separately to SLM
    - Combine summaries, don't embed chunks
    
    Args:
        text: Full text to split
        chunk_size: Target size of each chunk (characters)
        chunk_overlap: Overlap between consecutive chunks
    
    Returns:
        List of text chunks
    """
    text = clean_text(text)
    
    if chunk_size <= 0:
        return [text]
    
    chunks: list[str] = []
    start = 0
    length = len(text)
    
    # Ensure non-negative overlap smaller than chunk_size
    overlap = max(0, min(chunk_overlap, chunk_size - 1))
    
    while start < length:
        end = min(start + chunk_size, length)
        chunk = text[start:end].strip()
        
        if chunk:
            chunks.append(chunk)
        
        if end == length:
            break
        
        start = end - overlap
    
    return chunks


# ============================================================================
# SECTION 4: FILE UTILITIES
# ============================================================================

def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF file content.
    
    Args:
        content: PDF file content (bytes)
    
    Returns:
        Extracted text string
    """
    try:
        from PyPDF2 import PdfReader
        
        # Open PDF from bytes
        pdf_file = BytesIO(content)
        reader = PdfReader(pdf_file)
        
        # Extract text from all pages
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        # Combine and clean
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text)
    
    except ImportError:
        logger.warning("PyPDF2 not installed, returning placeholder")
        return "[PDF Content - PyPDF2 not available]"
    
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX (Microsoft Word) file content.
    
    Args:
        content: DOCX file content (bytes)
    
    Returns:
        Extracted text string
    """
    try:
        from docx import Document
        
        # Open DOCX from bytes
        doc_file = BytesIO(content)
        doc = Document(doc_file)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # Combine and clean
        full_text = "\n\n".join(text_parts)
        return clean_text(full_text)
    
    except ImportError:
        logger.warning("python-docx not installed, returning placeholder")
        return "[DOCX Content - python-docx not available]"
    
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text_from_file(
    content: bytes,
    filename: str,
    extension: str
) -> str:
    """
    Extract text from uploaded file (smart routing).
    
    ROUTING LOGIC:
    - .pdf → extract_text_from_pdf()
    - .docx → extract_text_from_docx()
    - .txt, .md → UTF-8 decode
    - Others → Try UTF-8 decode, fallback to Latin-1
    
    Args:
        content: File content (bytes)
        filename: Original filename
        extension: File extension (e.g., '.pdf', '.txt')
    
    Returns:
        Extracted text string
    
    Raises:
        ValueError: If extraction fails for all methods
    """
    logger.info(f"Extracting text from {filename} ({extension})")
    extension = extension.lower()
    
    try:
        if extension == '.pdf':
            return extract_text_from_pdf(content)
        
        elif extension == '.docx':
            return extract_text_from_docx(content)
        
        elif extension in ['.txt', '.md']:
            # Direct text decoding
            try:
                return clean_text(content.decode('utf-8'))
            except UnicodeDecodeError:
                # Fallback to other encodings
                try:
                    return clean_text(content.decode('latin-1'))
                except:
                    logger.warning(f"Cannot decode {filename}, using placeholder")
                    return f"[Binary file content - {filename}]"
        
        else:
            # Try UTF-8 decoding for unknown text files
            try:
                return clean_text(content.decode('utf-8'))
            except UnicodeDecodeError:
                logger.error(f"Cannot decode {extension} file")
                raise ValueError(f"Unsupported file format: {extension}")
    
    except ValueError as e:
        logger.error(f"Text extraction error: {str(e)}")
        raise
    
    except Exception as e:
        logger.error(f"Unexpected extraction error: {str(e)}")
        raise ValueError(f"Failed to extract text: {str(e)}")


def get_file_metadata(
    filename: str,
    content: bytes,
    content_hash: str,
    extension: str
) -> Dict[str, Any]:
    """
    Generate file metadata for logging and tracking.
    
    Args:
        filename: Original filename
        content: File content (bytes)
        content_hash: SHA256 hash of content
        extension: File extension
    
    Returns:
        Metadata dict
    """
    # Detect MIME type
    mime_type, encoding = mimetypes.guess_type(filename)
    mime_type = mime_type or 'application/octet-stream'
    
    return {
        'filename': filename,
        'extension': extension,
        'size_bytes': len(content),
        'content_hash': content_hash,
        'mime_type': mime_type,
        'encoding': encoding or 'binary'
    }
